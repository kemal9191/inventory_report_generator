import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import sys
import json


df1 = pd.read_excel("data/results.xlsx")

persons_raw_data= [] 
for index, row in df1.iterrows(): 
    person_raw_data = [row.Name, row.A, row.C, row.E, row.F, row.G, 
    row.H, row.I, row.L, row.M, row.O, row.Q1, row.Q3, 
    row.Q4, row.G1, row.G2, row.G3, row.G4, row.G5]
    persons_raw_data.append(person_raw_data)

df2 = pd.read_excel("data/report_data.xlsx")

report_data= [] 
for index, row in df2.iterrows(): 
    raw_data = [row.A, row.C, row.E, row.F, row.G, 
    row.H, row.I, row.L, row.M, row.O, row.Q1, row.Q3, 
    row.Q4, row.G1, row.G2, row.G3, row.G4, row.G5]
    report_data.append(raw_data)


with open("data/global_report_data.txt") as f:
    dict = json.load(f)    


subfactors = [1,2,3,4,5,6,7,8,9,10,11,12,13]


for person in persons_raw_data:
    doc = docx.Document()
    for index, item in enumerate(person):
        if index==0:
            name = item
            heading = f"{name} DBE Kişilik Envanteri Raporu".format(name)
            title = doc.add_heading(heading, 0)
            title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc_uri = f"reports/{name}.docx".format(name)
        if index in subfactors:
            factor_name = report_data[0][index-1]
            factor_title = f"{factor_name} ({item})".format(factor_name, item)
            header = doc.add_paragraph()
            header.add_run((factor_title)).bold = True
            header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if item < 36:
                para = doc.add_paragraph(report_data[1][index-1])
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.style.font.name = 'Times New Roman'

            if item > 35 and item < 66:
                para = doc.add_paragraph(report_data[2][index-1])
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.style.font.name = 'Times New Roman'
                
            if item > 65:
                para = doc.add_paragraph(report_data[3][index-1])
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.style.font.name = 'Times New Roman'

    for key, value in dict.items():
        factor_name = key
        result = person[report_data[0].index(key)+1]
        factor_title = f"{factor_name} ({result})"
        header = doc.add_paragraph()
        header.add_run((factor_title)).bold = True
        header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for k, val in value.items():
            if k == "default":
                if result<36:
                    para = doc.add_paragraph(val["düşük"])
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.style.font.name = 'Times New Roman'
                if result > 35 and result < 66:
                    para = doc.add_paragraph(val["orta"])                    
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.style.font.name = 'Times New Roman'
                if result > 65:
                    para = doc.add_paragraph(val["yüksek"])
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.style.font.name = 'Times New Roman'
            else:
                fresult = person[report_data[0].index(k)+1]
                if fresult < 36:
                    para = doc.add_paragraph(f"\t• {val['düşük']}")
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.style.font.name = 'Times New Roman'
                if fresult > 35 and fresult < 66:
                    para = doc.add_paragraph(f"\t• {val['orta']}")
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.style.font.name = 'Times New Roman'
                if fresult > 65:
                    para = doc.add_paragraph(f"\t• {val['yüksek']}")
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.style.font.name = 'Times New Roman'
    doc.save(doc_uri)